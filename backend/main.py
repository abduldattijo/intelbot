"""
Intelligence Document Analyzer Backend
FastAPI server for document analysis and intelligence extraction
"""

import os
import uuid
import time
import tempfile
from datetime import datetime
from typing import List, Dict, Any, Optional, Tuple
import logging

from fastapi import FastAPI, File, UploadFile, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel
import uvicorn

# NLP and Analysis Libraries
import spacy
import re
from collections import Counter, defaultdict
import nltk
from nltk.sentiment import SentimentIntensityAnalyzer
from nltk.tokenize import sent_tokenize, word_tokenize

# Document Processing Libraries
import PyPDF2
import docx
import pandas as pd
from io import StringIO, BytesIO

# Download required NLTK data
try:
    nltk.download('punkt', quiet=True)
    nltk.download('vader_lexicon', quiet=True)
    nltk.download('stopwords', quiet=True)
except:
    pass

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize FastAPI app
app = FastAPI(
    title="Intelligence Document Analyzer API",
    description="AI-Powered Security Intelligence Platform for Document Analysis",
    version="3.0.0"
)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://127.0.0.1:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Load spaCy model
try:
    nlp = spacy.load("en_core_web_sm")
    logger.info("SpaCy model loaded successfully")
except OSError:
    logger.error("SpaCy model not found. Please run: python -m spacy download en_core_web_sm")
    nlp = None

# Initialize sentiment analyzer
try:
    sia = SentimentIntensityAnalyzer()
    logger.info("NLTK sentiment analyzer loaded successfully")
except:
    logger.error("NLTK sentiment analyzer failed to load")
    sia = None


# Data Models
class DocumentMetadata(BaseModel):
    filename: str
    file_type: str
    uploaded_at: str
    file_size: int


class DocumentClassification(BaseModel):
    primary_type: str
    sub_types: List[str]
    confidence: float
    security_classification: str


class SentimentAnalysis(BaseModel):
    overall_sentiment: str
    sentiment_score: float
    threat_level: str
    urgency_indicators: List[str]


class GeographicIntelligence(BaseModel):
    states: List[str]
    cities: List[str]
    countries: List[str]
    coordinates: List[Dict[str, Any]]
    total_locations: int
    other_locations: List[str]


class TemporalIntelligence(BaseModel):
    dates_mentioned: List[str]
    time_periods: List[str]
    months_mentioned: List[str]
    years_mentioned: List[str]
    temporal_patterns: List[str]


class NumericalIntelligence(BaseModel):
    incidents: List[int]
    casualties: List[int]
    weapons: List[int]
    arrests: List[int]
    monetary_values: List[float]


class CrimePatterns(BaseModel):
    primary_crimes: List[Tuple[str, int]]
    crime_frequency: Dict[str, int]
    crime_trends: List[Dict[str, Any]]


class TextStatistics(BaseModel):
    word_count: int
    sentence_count: int
    paragraph_count: int
    readability_score: float
    language: str


class DocumentAnalysis(BaseModel):
    document_classification: DocumentClassification
    entities: Dict[str, List[str]]
    sentiment_analysis: SentimentAnalysis
    geographic_intelligence: GeographicIntelligence
    temporal_intelligence: TemporalIntelligence
    numerical_intelligence: NumericalIntelligence
    crime_patterns: CrimePatterns
    relationships: List[Dict[str, Any]]
    text_statistics: TextStatistics
    intelligence_summary: str
    confidence_score: float
    processing_time: float


class AnalyzedDocument(BaseModel):
    id: str
    content: str
    metadata: DocumentMetadata
    analysis: DocumentAnalysis


# Intelligence Analysis Engine
class IntelligenceAnalyzer:
    def __init__(self):
        self.crime_keywords = {
            'terrorism': ['terror', 'terrorist', 'bomb', 'explosive', 'attack', 'jihad', 'isis', 'al-qaeda'],
            'drug_trafficking': ['drug', 'cocaine', 'heroin', 'trafficking', 'smuggling', 'cartel', 'narcotic'],
            'human_trafficking': ['trafficking', 'slavery', 'prostitution', 'exploitation', 'forced labor'],
            'cybercrime': ['cyber', 'hacking', 'malware', 'phishing', 'ransomware', 'breach'],
            'organized_crime': ['mafia', 'gang', 'cartel', 'syndicate', 'racketeering'],
            'violence': ['murder', 'assault', 'kidnapping', 'robbery', 'violence', 'shooting'],
            'fraud': ['fraud', 'scam', 'embezzlement', 'money laundering', 'corruption'],
            'banditry': ['bandit', 'bandits', 'banditry', 'rustling', 'cattle rustling'],
            'armed_robbery': ['armed robbery', 'robbery', 'robber', 'robbers', 'armed robber']
        }

        self.weapon_keywords = [
            'gun', 'rifle', 'pistol', 'weapon', 'firearm', 'ammunition', 'explosive',
            'bomb', 'grenade', 'ak-47', 'ar-15', 'shotgun', 'revolver', 'ak 47',
            'pump action', 'locally made', 'double barrel'
        ]

        self.threat_indicators = [
            'threat', 'danger', 'attack', 'kill', 'destroy', 'eliminate', 'target',
            'strike', 'hit', 'bomb', 'explode', 'urgent', 'immediate', 'critical',
            'high threat', 'security threat', 'criminal activity'
        ]

        # Nigerian states and common locations (enhanced database)
        self.nigerian_states = [
            'lagos', 'kano', 'kaduna', 'oyo', 'rivers', 'bayelsa', 'cross river',
            'akwa ibom', 'abia', 'anambra', 'imo', 'enugu', 'ebonyi', 'delta',
            'edo', 'ondo', 'ekiti', 'osun', 'ogun', 'kwara', 'kogi', 'benue',
            'plateau', 'nasarawa', 'taraba', 'adamawa', 'borno', 'yobe', 'bauchi',
            'gombe', 'jigawa', 'katsina', 'kebbi', 'sokoto', 'zamfara', 'niger',
            'abuja', 'fct'
        ]

        # CORRECTED: Comprehensive Nigerian locations database with accurate coordinates
        self.nigerian_locations = {
            # Nigerian States mapped to their capitals (FIXED COORDINATES)
            'abia': {'lat': 5.5265, 'lng': 7.4906, 'type': 'state_capital', 'capital': 'Umuahia'},
            'adamawa': {'lat': 9.2000, 'lng': 12.4833, 'type': 'state_capital', 'capital': 'Yola'},
            'akwa ibom': {'lat': 5.0515, 'lng': 7.9307, 'type': 'state_capital', 'capital': 'Uyo'},
            'anambra': {'lat': 6.2120, 'lng': 7.0740, 'type': 'state_capital', 'capital': 'Awka'},
            'bauchi': {'lat': 10.3158, 'lng': 9.8442, 'type': 'state_capital', 'capital': 'Bauchi'},
            'bayelsa': {'lat': 4.9267, 'lng': 6.2676, 'type': 'state_capital', 'capital': 'Yenagoa'},
            'benue': {'lat': 7.7340, 'lng': 8.5120, 'type': 'state_capital', 'capital': 'Makurdi'},
            'borno': {'lat': 11.8311, 'lng': 13.1510, 'type': 'state_capital', 'capital': 'Maiduguri'},
            'cross river': {'lat': 4.9516, 'lng': 8.3220, 'type': 'state_capital', 'capital': 'Calabar'},
            'delta': {'lat': 6.1677, 'lng': 6.7337, 'type': 'state_capital', 'capital': 'Asaba'},
            'ebonyi': {'lat': 6.3248, 'lng': 8.1142, 'type': 'state_capital', 'capital': 'Abakaliki'},
            'edo': {'lat': 6.3350, 'lng': 5.6037, 'type': 'state_capital', 'capital': 'Benin City'},
            'ekiti': {'lat': 7.6667, 'lng': 5.2167, 'type': 'state_capital', 'capital': 'Ado-Ekiti'},
            'enugu': {'lat': 6.5244, 'lng': 7.5112, 'type': 'state_capital', 'capital': 'Enugu'},
            'gombe': {'lat': 10.2840, 'lng': 11.1610, 'type': 'state_capital', 'capital': 'Gombe'},
            'imo': {'lat': 5.4840, 'lng': 7.0351, 'type': 'state_capital', 'capital': 'Owerri'},  # CORRECTED
            'jigawa': {'lat': 11.7564, 'lng': 9.3388, 'type': 'state_capital', 'capital': 'Dutse'},
            'kaduna': {'lat': 10.5105, 'lng': 7.4165, 'type': 'state_capital', 'capital': 'Kaduna'},  # CORRECTED
            'kano': {'lat': 12.0022, 'lng': 8.5920, 'type': 'state_capital', 'capital': 'Kano'},
            'katsina': {'lat': 12.9908, 'lng': 7.6018, 'type': 'state_capital', 'capital': 'Katsina'},
            'kebbi': {'lat': 12.4537, 'lng': 4.1994, 'type': 'state_capital', 'capital': 'Birnin Kebbi'},
            'kogi': {'lat': 7.7974, 'lng': 6.7337, 'type': 'state_capital', 'capital': 'Lokoja'},
            'kwara': {'lat': 8.5000, 'lng': 4.5500, 'type': 'state_capital', 'capital': 'Ilorin'},
            'lagos': {'lat': 6.5962, 'lng': 3.3431, 'type': 'state_capital', 'capital': 'Ikeja'},
            'nasarawa': {'lat': 8.4833, 'lng': 8.5167, 'type': 'state_capital', 'capital': 'Lafia'},
            'niger': {'lat': 9.6134, 'lng': 6.5560, 'type': 'state_capital', 'capital': 'Minna'},
            'ogun': {'lat': 7.1475, 'lng': 3.3619, 'type': 'state_capital', 'capital': 'Abeokuta'},
            'ondo': {'lat': 7.2571, 'lng': 5.2058, 'type': 'state_capital', 'capital': 'Akure'},
            'osun': {'lat': 7.7719, 'lng': 4.5567, 'type': 'state_capital', 'capital': 'Oshogbo'},
            'oyo': {'lat': 7.3775, 'lng': 3.9470, 'type': 'state_capital', 'capital': 'Ibadan'},
            'plateau': {'lat': 9.8965, 'lng': 8.8583, 'type': 'state_capital', 'capital': 'Jos'},
            'rivers': {'lat': 4.8156, 'lng': 7.0498, 'type': 'state_capital', 'capital': 'Port Harcourt'},
            'sokoto': {'lat': 13.0609, 'lng': 5.2476, 'type': 'state_capital', 'capital': 'Sokoto'},
            'taraba': {'lat': 8.8833, 'lng': 11.3667, 'type': 'state_capital', 'capital': 'Jalingo'},  # CORRECTED
            'yobe': {'lat': 11.7469, 'lng': 11.9609, 'type': 'state_capital', 'capital': 'Damaturu'},
            'zamfara': {'lat': 12.1667, 'lng': 6.6611, 'type': 'state_capital', 'capital': 'Gusau'},
            'abuja': {'lat': 9.0765, 'lng': 7.3986, 'type': 'federal_capital', 'capital': 'Abuja'},
            'fct': {'lat': 9.0765, 'lng': 7.3986, 'type': 'federal_capital', 'capital': 'Abuja'},

            # Major cities and state capitals by city name
            'umuahia': {'lat': 5.5265, 'lng': 7.4906, 'type': 'state_capital'},
            'yola': {'lat': 9.2000, 'lng': 12.4833, 'type': 'state_capital'},
            'uyo': {'lat': 5.0515, 'lng': 7.9307, 'type': 'state_capital'},
            'awka': {'lat': 6.2120, 'lng': 7.0740, 'type': 'state_capital'},
            'bauchi': {'lat': 10.3158, 'lng': 9.8442, 'type': 'state_capital'},
            'yenagoa': {'lat': 4.9267, 'lng': 6.2676, 'type': 'state_capital'},
            'makurdi': {'lat': 7.7340, 'lng': 8.5120, 'type': 'state_capital'},
            'maiduguri': {'lat': 11.8311, 'lng': 13.1510, 'type': 'state_capital'},
            'calabar': {'lat': 4.9516, 'lng': 8.3220, 'type': 'state_capital'},
            'asaba': {'lat': 6.1677, 'lng': 6.7337, 'type': 'state_capital'},
            'abakaliki': {'lat': 6.3248, 'lng': 8.1142, 'type': 'state_capital'},
            'benin city': {'lat': 6.3350, 'lng': 5.6037, 'type': 'state_capital'},
            'ado-ekiti': {'lat': 7.6667, 'lng': 5.2167, 'type': 'state_capital'},
            'enugu': {'lat': 6.5244, 'lng': 7.5112, 'type': 'state_capital'},
            'gombe': {'lat': 10.2840, 'lng': 11.1610, 'type': 'state_capital'},
            'owerri': {'lat': 5.4840, 'lng': 7.0351, 'type': 'state_capital'},
            'dutse': {'lat': 11.7564, 'lng': 9.3388, 'type': 'state_capital'},
            'kaduna': {'lat': 10.5105, 'lng': 7.4165, 'type': 'state_capital'},
            'kano': {'lat': 12.0022, 'lng': 8.5920, 'type': 'state_capital'},
            'katsina': {'lat': 12.9908, 'lng': 7.6018, 'type': 'state_capital'},
            'birnin kebbi': {'lat': 12.4537, 'lng': 4.1994, 'type': 'state_capital'},
            'lokoja': {'lat': 7.7974, 'lng': 6.7337, 'type': 'state_capital'},
            'ilorin': {'lat': 8.5000, 'lng': 4.5500, 'type': 'state_capital'},
            'ikeja': {'lat': 6.5962, 'lng': 3.3431, 'type': 'state_capital'},
            'lafia': {'lat': 8.4833, 'lng': 8.5167, 'type': 'state_capital'},
            'minna': {'lat': 9.6134, 'lng': 6.5560, 'type': 'state_capital'},
            'abeokuta': {'lat': 7.1475, 'lng': 3.3619, 'type': 'state_capital'},
            'akure': {'lat': 7.2571, 'lng': 5.2058, 'type': 'state_capital'},
            'oshogbo': {'lat': 7.7719, 'lng': 4.5567, 'type': 'state_capital'},
            'ibadan': {'lat': 7.3775, 'lng': 3.9470, 'type': 'state_capital'},
            'jos': {'lat': 9.8965, 'lng': 8.8583, 'type': 'state_capital'},
            'port harcourt': {'lat': 4.8156, 'lng': 7.0498, 'type': 'state_capital'},
            'sokoto': {'lat': 13.0609, 'lng': 5.2476, 'type': 'state_capital'},
            'jalingo': {'lat': 8.8833, 'lng': 11.3667, 'type': 'state_capital'},
            'damaturu': {'lat': 11.7469, 'lng': 11.9609, 'type': 'state_capital'},
            'gusau': {'lat': 12.1667, 'lng': 6.6611, 'type': 'state_capital'},

            # Major commercial cities and towns
            'lagos city': {'lat': 6.5244, 'lng': 3.3792, 'type': 'major_city'},
            'warri': {'lat': 5.5167, 'lng': 5.7500, 'type': 'major_city'},
            'aba': {'lat': 5.1068, 'lng': 7.3668, 'type': 'major_city'},
            'onitsha': {'lat': 6.1667, 'lng': 6.7833, 'type': 'major_city'},
            'zaria': {'lat': 11.0804, 'lng': 7.7170, 'type': 'major_city'},
            'ife': {'lat': 7.4905, 'lng': 4.5621, 'type': 'major_city'},
            'ogbomoso': {'lat': 8.1336, 'lng': 4.2570, 'type': 'major_city'},
            'sapele': {'lat': 5.8939, 'lng': 5.6760, 'type': 'major_city'},
            'okene': {'lat': 7.5519, 'lng': 6.2350, 'type': 'major_city'},

            # Important LGAs and districts
            'victoria island': {'lat': 6.4281, 'lng': 3.4219, 'type': 'district'},
            'apapa': {'lat': 6.4474, 'lng': 3.3903, 'type': 'commercial'},
            'ikoyi': {'lat': 6.4525, 'lng': 3.4328, 'type': 'district'},
            'surulere': {'lat': 6.5027, 'lng': 3.3584, 'type': 'lga'},
            'alimosho': {'lat': 6.5833, 'lng': 3.2667, 'type': 'lga'},
            'agege': {'lat': 6.6186, 'lng': 3.3403, 'type': 'lga'},
            'ikorodu': {'lat': 6.6019, 'lng': 3.5106, 'type': 'lga'},
            'epe': {'lat': 6.5833, 'lng': 3.9833, 'type': 'lga'},
            'badagry': {'lat': 6.4319, 'lng': 2.8878, 'type': 'lga'},

            # Border towns and strategic locations
            'bonny': {'lat': 4.4500, 'lng': 7.1667, 'type': 'oil_terminal'},
            'forcados': {'lat': 5.3833, 'lng': 5.4000, 'type': 'oil_terminal'},
            'escravos': {'lat': 5.5333, 'lng': 5.0667, 'type': 'oil_terminal'},
            'brass': {'lat': 4.3167, 'lng': 6.2333, 'type': 'oil_terminal'},

            # Airports
            'murtala muhammed airport': {'lat': 6.5772, 'lng': 3.3211, 'type': 'airport'},
            'nnamdi azikiwe airport': {'lat': 9.0063, 'lng': 7.2631, 'type': 'airport'},
            'mallam aminu kano airport': {'lat': 12.0476, 'lng': 8.5246, 'type': 'airport'},
            'port harcourt airport': {'lat': 5.0156, 'lng': 6.9496, 'type': 'airport'},
        }

    def extract_text_from_file(self, file_content: bytes, filename: str) -> str:
        """Extract text content from various file formats"""
        try:
            file_extension = filename.lower().split('.')[-1]

            if file_extension == 'pdf':
                return self._extract_from_pdf(file_content)
            elif file_extension in ['doc', 'docx']:
                return self._extract_from_docx(file_content)
            elif file_extension == 'txt':
                return file_content.decode('utf-8', errors='ignore')
            elif file_extension in ['csv', 'xlsx', 'xls']:
                return self._extract_from_excel(file_content)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")

        except Exception as e:
            logger.error(f"Text extraction error: {str(e)}")
            raise HTTPException(status_code=400, detail=f"Failed to extract text: {str(e)}")

    def _extract_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF files"""
        try:
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            raise ValueError(f"PDF extraction failed: {str(e)}")

    def _extract_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX files"""
        try:
            doc = docx.Document(BytesIO(file_content))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            raise ValueError(f"DOCX extraction failed: {str(e)}")

    def _extract_from_excel(self, file_content: bytes) -> str:
        """Extract text from Excel/CSV files"""
        try:
            # Try reading as Excel first
            try:
                df = pd.read_excel(BytesIO(file_content))
            except:
                # Fallback to CSV
                df = pd.read_csv(StringIO(file_content.decode('utf-8', errors='ignore')))

            # Convert DataFrame to text
            text = df.to_string()
            return text
        except Exception as e:
            raise ValueError(f"Excel/CSV extraction failed: {str(e)}")

    def analyze_document(self, text: str, metadata: DocumentMetadata) -> DocumentAnalysis:
        """Perform comprehensive intelligence analysis on document text"""
        start_time = time.time()

        # Basic text preprocessing
        text_lower = text.lower()
        sentences = sent_tokenize(text)
        words = word_tokenize(text_lower)

        # Extract entities using spaCy
        entities = self._extract_entities(text)

        # Analyze sentiment and threat level
        sentiment_analysis = self._analyze_sentiment(text, text_lower)

        # Extract geographic intelligence
        geographic_intel = self._extract_geographic_intelligence(text, text_lower)

        # Extract temporal intelligence
        temporal_intel = self._extract_temporal_intelligence(text)

        # Extract numerical intelligence
        numerical_intel = self._extract_numerical_intelligence(text)

        # Analyze crime patterns
        crime_patterns = self._analyze_crime_patterns(text_lower)

        # Classify document
        doc_classification = self._classify_document(text_lower, crime_patterns)

        # Extract relationships
        relationships = self._extract_relationships(text)

        # Calculate text statistics
        text_stats = self._calculate_text_statistics(text, sentences, words)

        # Generate intelligence summary (PROTECTED - ensure it doesn't get deleted)
        intelligence_summary = self._generate_intelligence_summary(
            entities, sentiment_analysis, geographic_intel, crime_patterns, text_stats
        )

        # Calculate overall confidence score
        confidence_score = self._calculate_confidence_score(
            entities, sentiment_analysis, geographic_intel, temporal_intel
        )

        processing_time = time.time() - start_time

        return DocumentAnalysis(
            document_classification=doc_classification,
            entities=entities,
            sentiment_analysis=sentiment_analysis,
            geographic_intelligence=geographic_intel,
            temporal_intelligence=temporal_intel,
            numerical_intelligence=numerical_intel,
            crime_patterns=crime_patterns,
            relationships=relationships,
            text_statistics=text_stats,
            intelligence_summary=intelligence_summary,  # PROTECTED
            confidence_score=confidence_score,
            processing_time=processing_time
        )

    def _extract_entities(self, text: str) -> Dict[str, List[str]]:
        """Extract named entities using spaCy NLP"""
        entities = {
            'persons': [],
            'organizations': [],
            'locations': [],
            'weapons': [],
            'vehicles': [],
            'dates': []
        }

        if nlp is None:
            return entities

        try:
            doc = nlp(text)

            for ent in doc.ents:
                if ent.label_ == "PERSON":
                    entities['persons'].append(ent.text)
                elif ent.label_ == "ORG":
                    entities['organizations'].append(ent.text)
                elif ent.label_ in ["GPE", "LOC"]:
                    entities['locations'].append(ent.text)
                elif ent.label_ == "DATE":
                    entities['dates'].append(ent.text)

            # Extract weapons using keyword matching
            text_lower = text.lower()
            for weapon in self.weapon_keywords:
                if weapon in text_lower:
                    entities['weapons'].append(weapon)

            # Extract vehicles using simple patterns
            vehicle_patterns = [
                r'\b(car|truck|motorcycle|bike|vehicle|van|suv)\b',
                r'\b(toyota|honda|ford|mercedes|bmw|volkswagen)\b',
                r'\b(boxer|motorcycles)\b'
            ]
            for pattern in vehicle_patterns:
                matches = re.findall(pattern, text_lower)
                entities['vehicles'].extend(matches)

            # Remove duplicates and clean up
            for key in entities:
                entities[key] = list(set(entities[key]))

        except Exception as e:
            logger.error(f"Entity extraction error: {str(e)}")

        return entities

    def _analyze_sentiment(self, text: str, text_lower: str) -> SentimentAnalysis:
        """Analyze sentiment and determine threat level"""
        try:
            if sia:
                scores = sia.polarity_scores(text)
                sentiment_score = scores['compound']

                if sentiment_score >= 0.05:
                    overall_sentiment = "positive"
                elif sentiment_score <= -0.05:
                    overall_sentiment = "negative"
                else:
                    overall_sentiment = "neutral"
            else:
                # Fallback simple sentiment analysis
                positive_words = ['good', 'great', 'excellent', 'positive', 'success', 'decrease', 'reduced']
                negative_words = ['bad', 'terrible', 'negative', 'fail', 'problem', 'increase', 'surge', 'criminal']

                pos_count = sum(1 for word in positive_words if word in text_lower)
                neg_count = sum(1 for word in negative_words if word in text_lower)

                if pos_count > neg_count:
                    overall_sentiment = "positive"
                    sentiment_score = 0.1
                elif neg_count > pos_count:
                    overall_sentiment = "negative"
                    sentiment_score = -0.1
                else:
                    overall_sentiment = "neutral"
                    sentiment_score = 0.0

            # Determine threat level based on keywords and crime indicators
            threat_count = sum(1 for indicator in self.threat_indicators if indicator in text_lower)

            # Add specific crime-related threat indicators
            crime_indicators = ['armed robbery', 'murder', 'kidnapping', 'banditry', 'cattle rustling', 'terrorism']
            crime_count = sum(1 for crime in crime_indicators if crime in text_lower)

            total_threat_score = threat_count + (crime_count * 2)  # Weight crime indicators more

            if total_threat_score >= 8:
                threat_level = "High"
            elif total_threat_score >= 4:
                threat_level = "Medium"
            else:
                threat_level = "Low"

            # Extract urgency indicators
            urgency_indicators = [indicator for indicator in self.threat_indicators if indicator in text_lower]
            urgency_indicators.extend([crime for crime in crime_indicators if crime in text_lower])

            return SentimentAnalysis(
                overall_sentiment=overall_sentiment,
                sentiment_score=sentiment_score,
                threat_level=threat_level,
                urgency_indicators=list(set(urgency_indicators))[:10]  # Remove duplicates and limit to 10
            )

        except Exception as e:
            logger.error(f"Sentiment analysis error: {str(e)}")
            return SentimentAnalysis(
                overall_sentiment="neutral",
                sentiment_score=0.0,
                threat_level="Low",
                urgency_indicators=[]
            )

    def _extract_geographic_intelligence(self, text: str, text_lower: str) -> GeographicIntelligence:
        """Extract geographic information and locations with FIXED coordinate mapping"""
        try:
            states = []
            cities = []
            countries = []
            coordinates = []
            other_locations = []

            # Extract Nigerian states (exact matching only)
            for state in self.nigerian_states:
                # Check for exact matches and common variations
                variations = [state, state.replace(' ', ''), f"{state} state"]
                for variation in variations:
                    if f" {variation} " in f" {text_lower} " or variation == text_lower:
                        state_title = state.title()
                        if state_title not in states:
                            states.append(state_title)

            # Enhanced coordinate extraction patterns
            coord_patterns = [
                # Decimal degrees: lat, lon (most common format)
                r'(\d+\.?\d*)[°\s,]*\s*[-,]\s*(\d+\.?\d*)[°\s]*',
                # Coordinates with explicit lat/lon labels
                r'lat[itude]*[:=\s]*(\d+\.?\d*)[,\s]*lon[gitude]*[:=\s]*(\d+\.?\d*)',
                # Nigerian format: N, E coordinates
                r'(\d+\.?\d*)[°\s]*N[orth]*[,\s]*(\d+\.?\d*)[°\s]*E[ast]*',
                # GPS coordinates in brackets or parentheses
                r'[\[\(](\d+\.?\d*)[,\s]+(\d+\.?\d*)[\]\)]',
                # Coordinates word "coordinates" nearby
                r'coordinates[:\s]*(\d+\.?\d*)[,\s]+(\d+\.?\d*)',
                # Simple decimal format with comma separation
                r'\b(\d+\.\d{4}),\s*(\d+\.\d{4})\b',
                # Position/located at format
                r'(?:position|located at)[:\s]*(\d+\.?\d*)[,\s]+(\d+\.?\d*)'
            ]

            coordinate_matches_found = []

            for pattern in coord_patterns:
                matches = re.finditer(pattern, text, re.IGNORECASE)
                for match in matches:
                    try:
                        groups = match.groups()
                        if len(groups) >= 2:
                            lat, lon = float(groups[0]), float(groups[1])

                            # Validate coordinates for Nigeria region (extended bounds)
                            if 3.0 <= lat <= 15.0 and 2.0 <= lon <= 16.0:
                                # Find nearby location name
                                location_name = self._find_location_context(text, match.start(), match.end())

                                coord_info = {
                                    'latitude': round(lat, 6),
                                    'longitude': round(lon, 6),
                                    'location_name': location_name or f'Coordinates {lat:.4f}, {lon:.4f}',
                                    'confidence': 0.9
                                }

                                # Avoid duplicates
                                is_duplicate = any(
                                    abs(coord['latitude'] - coord_info['latitude']) < 0.001 and
                                    abs(coord['longitude'] - coord_info['longitude']) < 0.001
                                    for coord in coordinate_matches_found
                                )

                                if not is_duplicate:
                                    coordinate_matches_found.append(coord_info)

                    except (ValueError, IndexError):
                        continue

            coordinates = coordinate_matches_found

            # Extract locations using spaCy if available with EXACT matching
            if nlp:
                doc = nlp(text)
                for ent in doc.ents:
                    if ent.label_ == "GPE":  # Geopolitical entity
                        location = ent.text.strip()
                        location_lower = location.lower()

                        # Check if it's a Nigerian state using exact mapping
                        if location_lower in self.nigerian_locations:
                            if location_lower in [s.lower() for s in self.nigerian_states]:
                                if location.title() not in states:
                                    states.append(location.title())
                            else:
                                if location.title() not in cities:
                                    cities.append(location.title())
                        # Check if it's a country
                        elif location_lower in ['nigeria', 'cameroon', 'chad', 'niger', 'benin']:
                            if location.title() not in countries:
                                countries.append(location.title())
                        else:
                            # Add to other locations if not already categorized and longer than 2 chars
                            if (location not in other_locations and
                                    len(location) > 2 and
                                    location_lower not in ['the', 'and', 'of', 'in', 'at']):
                                other_locations.append(location)

            # FIXED: Map detected locations to coordinates using EXACT matching only
            all_detected_locations = states + cities + other_locations

            for location in all_detected_locations:
                location_lower = location.lower().strip()

                # Remove common suffixes to normalize
                cleaned_location = location_lower.replace(' state', '').replace(' lga', '').strip()

                # Try exact match first (cleaned)
                location_data = None
                if cleaned_location in self.nigerian_locations:
                    location_data = self.nigerian_locations[cleaned_location]
                # Try original name
                elif location_lower in self.nigerian_locations:
                    location_data = self.nigerian_locations[location_lower]

                if location_data:
                    display_name = location
                    if 'capital' in location_data:
                        display_name = f"{location} (Capital: {location_data['capital']})"

                    coord_info = {
                        'latitude': round(location_data['lat'], 6),
                        'longitude': round(location_data['lng'], 6),
                        'location_name': display_name,
                        'confidence': 0.95  # High confidence for exact matches
                    }

                    # Check for duplicates
                    is_duplicate = any(
                        abs(coord['latitude'] - coord_info['latitude']) < 0.001 and
                        abs(coord['longitude'] - coord_info['longitude']) < 0.001
                        for coord in coordinates
                    )

                    if not is_duplicate:
                        coordinates.append(coord_info)

            # Remove duplicates from lists
            states = list(set(states))
            cities = list(set(cities))
            countries = list(set(countries))
            other_locations = list(set(other_locations))

            # Remove overlaps (don't include states in other_locations)
            other_locations = [loc for loc in other_locations if
                               not any(state.lower() in loc.lower() for state in states)]

            total_locations = len(states) + len(cities) + len(countries) + len(other_locations) + len(coordinates)

            return GeographicIntelligence(
                states=states,
                cities=cities,
                countries=countries,
                coordinates=coordinates,
                total_locations=total_locations,
                other_locations=other_locations[:20]  # Limit to 20 to avoid clutter
            )

        except Exception as e:
            logger.error(f"Geographic intelligence error: {str(e)}")
            return GeographicIntelligence(
                states=[], cities=[], countries=[], coordinates=[],
                total_locations=0, other_locations=[]
            )

    def _find_location_context(self, text: str, start: int, end: int) -> Optional[str]:
        """Find location name near coordinates"""
        try:
            # Look for location names within 100 characters before or after coordinates
            context_start = max(0, start - 100)
            context_end = min(len(text), end + 100)
            context = text[context_start:context_end]

            # Extract potential location names
            location_patterns = [
                r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b',  # Capitalized words
                r'\b(?:near|at|in|around)\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\b'  # Near/at/in location
            ]

            for pattern in location_patterns:
                matches = re.findall(pattern, context)
                if matches:
                    # Return the first reasonable match
                    for match in matches:
                        if isinstance(match, tuple):
                            match = match[0] if match[0] else match[1]
                        if len(match) > 2 and not match.lower() in ['the', 'and', 'of', 'in', 'at']:
                            return match

            return None
        except Exception:
            return None

    def _extract_temporal_intelligence(self, text: str) -> TemporalIntelligence:
        """Extract temporal information and patterns"""
        try:
            dates_mentioned = []
            time_periods = []
            months_mentioned = []
            years_mentioned = []
            temporal_patterns = []

            # Extract dates using regex patterns
            date_patterns = [
                r'\b\d{1,2}[/-]\d{1,2}[/-]\d{4}\b',  # MM/DD/YYYY or DD/MM/YYYY
                r'\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b',  # YYYY/MM/DD
                r'\b\w+ \d{1,2}, \d{4}\b',  # Month DD, YYYY
                r'\b\d{1,2} \w+ \d{4}\b',  # DD Month YYYY
                r'\b\w+ \d{4}\b'  # Month YYYY
            ]

            for pattern in date_patterns:
                matches = re.findall(pattern, text)
                dates_mentioned.extend(matches)

            # Extract months
            months = ['january', 'february', 'march', 'april', 'may', 'june',
                      'july', 'august', 'september', 'october', 'november', 'december']
            text_lower = text.lower()

            for month in months:
                if month in text_lower:
                    months_mentioned.append(month.title())

            # Extract years (2019-2025 range for recent intelligence)
            year_pattern = r'\b(20[12][0-9])\b'
            years = re.findall(year_pattern, text)
            years_mentioned = list(set(years))

            # Extract time periods
            time_period_keywords = ['morning', 'afternoon', 'evening', 'night', 'dawn', 'dusk']
            for period in time_period_keywords:
                if period in text_lower:
                    time_periods.append(period)

            # Identify temporal patterns
            if len(dates_mentioned) > 1:
                temporal_patterns.append("Multiple dates referenced")
            if len(years_mentioned) > 1:
                temporal_patterns.append("Multi-year timeline")
            if len(months_mentioned) > 3:
                temporal_patterns.append("Extended period analysis")

            return TemporalIntelligence(
                dates_mentioned=list(set(dates_mentioned)),
                time_periods=list(set(time_periods)),
                months_mentioned=list(set(months_mentioned)),
                years_mentioned=years_mentioned,
                temporal_patterns=temporal_patterns
            )

        except Exception as e:
            logger.error(f"Temporal intelligence error: {str(e)}")
            return TemporalIntelligence(
                dates_mentioned=[], time_periods=[], months_mentioned=[],
                years_mentioned=[], temporal_patterns=[]
            )

    def _extract_numerical_intelligence(self, text: str) -> NumericalIntelligence:
        """Extract numerical data and statistics"""
        try:
            incidents = []
            casualties = []
            weapons = []
            arrests = []
            monetary_values = []

            # Extract numbers using regex
            number_pattern = r'\b\d+\b'
            numbers = [int(match) for match in re.findall(number_pattern, text)]

            # Context-based classification with improved patterns
            text_lower = text.lower()

            # Extract incident counts
            incident_patterns = [
                r'(\d+)\s*(incident|attack|occurrence|case)',
                r'(incident|attack|occurrence|case).*?(\d+)',
                r'total.*?(\d+).*?(incident|case)',
                r'(\d+).*?(criminal activities|crime)'
            ]
            for pattern in incident_patterns:
                matches = re.findall(pattern, text_lower)
                for match in matches:
                    for item in match:
                        if item.isdigit():
                            incidents.append(int(item))

            # Extract casualty counts
            casualty_patterns = [
                r'(\d+)\s*(dead|killed|casualt|victim|injur|died|death)',
                r'(dead|killed|casualt|victim|injur|died|death).*?(\d+)',
                r'(\d+).*?person.*?(lost their lives|died)',
                r'about\s+(\d+)\s+persons'
            ]
            for pattern in casualty_patterns:
                matches = re.findall(pattern, text_lower)
                for match in matches:
                    for item in match:
                        if item.isdigit():
                            casualties.append(int(item))

            # Extract weapon counts
            weapon_patterns = [
                r'(\d+)\s*(gun|rifle|weapon|firearm|ak|ammunition)',
                r'(gun|rifle|weapon|firearm|ak|ammunition).*?(\d+)',
                r'(\d+).*?(round|cartridge|bullet)'
            ]
            for pattern in weapon_patterns:
                matches = re.findall(pattern, text_lower)
                for match in matches:
                    for item in match:
                        if item.isdigit():
                            weapons.append(int(item))

            # Extract arrest counts
            arrest_patterns = [
                r'(\d+)\s*(arrest|detain|capture|suspect)',
                r'(arrest|detain|capture|suspect).*?(\d+)',
                r'about\s+(\d+)\s+suspect'
            ]
            for pattern in arrest_patterns:
                matches = re.findall(pattern, text_lower)
                for match in matches:
                    for item in match:
                        if item.isdigit():
                            arrests.append(int(item))

            # Extract monetary values
            money_pattern = r'[$₦]\s*(\d+(?:,\d{3})*(?:\.\d{2})?)'
            money_matches = re.findall(money_pattern, text)
            for match in money_matches:
                try:
                    value = float(match.replace(',', ''))
                    monetary_values.append(value)
                except ValueError:
                    continue

            return NumericalIntelligence(
                incidents=incidents,
                casualties=casualties,
                weapons=weapons,
                arrests=arrests,
                monetary_values=monetary_values
            )

        except Exception as e:
            logger.error(f"Numerical intelligence error: {str(e)}")
            return NumericalIntelligence(
                incidents=[], casualties=[], weapons=[],
                arrests=[], monetary_values=[]
            )

    def _analyze_crime_patterns(self, text_lower: str) -> CrimePatterns:
        """Analyze crime patterns and frequencies"""
        try:
            crime_frequency = defaultdict(int)

            # Count crime-related keywords with enhanced detection
            for crime_type, keywords in self.crime_keywords.items():
                for keyword in keywords:
                    if keyword in text_lower:
                        crime_frequency[crime_type] += text_lower.count(keyword)

            # Get primary crimes (top 5)
            primary_crimes = sorted(crime_frequency.items(), key=lambda x: x[1], reverse=True)[:5]

            # Generate crime trends (simplified)
            crime_trends = []
            for crime_type, count in primary_crimes:
                if count > 5:
                    trend = "increasing"
                elif count > 2:
                    trend = "stable"
                else:
                    trend = "decreasing"

                crime_trends.append({
                    'crime_type': crime_type.replace('_', ' ').title(),
                    'trend': trend,
                    'confidence': min(count / 10.0, 1.0)
                })

            return CrimePatterns(
                primary_crimes=primary_crimes,
                crime_frequency=dict(crime_frequency),
                crime_trends=crime_trends
            )

        except Exception as e:
            logger.error(f"Crime pattern analysis error: {str(e)}")
            return CrimePatterns(
                primary_crimes=[], crime_frequency={}, crime_trends=[]
            )

    def _classify_document(self, text_lower: str, crime_patterns: CrimePatterns) -> DocumentClassification:
        """Classify the document type and security level"""
        try:
            # Determine primary type based on content
            if any(keyword in text_lower for keyword in ['report', 'incident', 'case', 'returns']):
                primary_type = "incident_report"
            elif any(keyword in text_lower for keyword in ['intelligence', 'briefing', 'analysis']):
                primary_type = "intelligence_briefing"
            elif any(keyword in text_lower for keyword in ['surveillance', 'monitoring', 'observation']):
                primary_type = "surveillance_report"
            elif any(keyword in text_lower for keyword in ['threat', 'warning', 'alert']):
                primary_type = "threat_assessment"
            elif any(keyword in text_lower for keyword in ['armed banditry', 'robbery', 'criminal activities']):
                primary_type = "security_report"
            else:
                primary_type = "general_document"

            # Determine sub-types based on crime patterns
            sub_types = []
            if crime_patterns.crime_frequency:
                top_crimes = sorted(crime_patterns.crime_frequency.items(),
                                    key=lambda x: x[1], reverse=True)[:3]
                sub_types = [crime.replace('_', ' ').title() for crime, count in top_crimes if count > 0]

            # Determine security classification
            sensitive_keywords = ['classified', 'confidential', 'secret', 'restricted']
            if any(keyword in text_lower for keyword in sensitive_keywords):
                security_classification = "CLASSIFIED"
            elif len(crime_patterns.primary_crimes) > 3:
                security_classification = "CONFIDENTIAL"
            elif any(keyword in text_lower for keyword in ['public', 'open', 'unclassified']):
                security_classification = "UNCLASSIFIED"
            else:
                security_classification = "RESTRICTED"

            # Calculate confidence based on keyword matches
            confidence = min(len(sub_types) / 3.0 + 0.3, 1.0)

            return DocumentClassification(
                primary_type=primary_type,
                sub_types=sub_types,
                confidence=confidence,
                security_classification=security_classification
            )

        except Exception as e:
            logger.error(f"Document classification error: {str(e)}")
            return DocumentClassification(
                primary_type="general_document",
                sub_types=[],
                confidence=0.5,
                security_classification="UNCLASSIFIED"
            )

    def _extract_relationships(self, text: str) -> List[Dict[str, Any]]:
        """Extract relationships between entities"""
        relationships = []

        try:
            if nlp is None:
                return relationships

            doc = nlp(text)

            # Simple relationship extraction based on sentence structure
            for sent in doc.sents:
                entities_in_sent = [ent for ent in sent.ents if ent.label_ in ["PERSON", "ORG", "GPE"]]

                if len(entities_in_sent) >= 2:
                    for i in range(len(entities_in_sent) - 1):
                        relationships.append({
                            'entity1': entities_in_sent[i].text,
                            'entity2': entities_in_sent[i + 1].text,
                            'relationship_type': 'mentioned_together',
                            'confidence': 0.6
                        })

                        # Limit to avoid too many relationships
                        if len(relationships) >= 20:
                            return relationships

        except Exception as e:
            logger.error(f"Relationship extraction error: {str(e)}")

        return relationships

    def _calculate_text_statistics(self, text: str, sentences: List[str], words: List[str]) -> TextStatistics:
        """Calculate basic text statistics"""
        try:
            word_count = len(words)
            sentence_count = len(sentences)
            paragraph_count = len([p for p in text.split('\n\n') if p.strip()])

            # Simple readability score (Flesch-like approximation)
            if sentence_count > 0 and word_count > 0:
                avg_sentence_length = word_count / sentence_count
                readability_score = max(0, min(100, 206.835 - (1.015 * avg_sentence_length)))
            else:
                readability_score = 50

            return TextStatistics(
                word_count=word_count,
                sentence_count=sentence_count,
                paragraph_count=paragraph_count,
                readability_score=readability_score,
                language="en"
            )

        except Exception as e:
            logger.error(f"Text statistics error: {str(e)}")
            return TextStatistics(
                word_count=0, sentence_count=0, paragraph_count=0,
                readability_score=0, language="en"
            )

    def _generate_intelligence_summary(self, entities, sentiment_analysis,
                                       geographic_intel, crime_patterns, text_stats=None) -> str:
        """Generate AI-powered intelligence summary (PROTECTED FROM DELETION)"""
        try:
            summary_parts = []

            # Document overview
            threat_level = sentiment_analysis.threat_level
            summary_parts.append(f"Intelligence analysis indicates a {threat_level.lower()} threat level.")

            # Entity summary
            total_entities = sum(len(ent_list) for ent_list in entities.values())
            if total_entities > 0:
                summary_parts.append(f"Document contains {total_entities} identified entities including "
                                     f"{len(entities['persons'])} persons, {len(entities['organizations'])} organizations, "
                                     f"and {len(entities['locations'])} locations.")

            # Geographic summary
            if geographic_intel.total_locations > 0:
                summary_parts.append(f"Geographic analysis identified {geographic_intel.total_locations} locations "
                                     f"with {len(geographic_intel.states)} Nigerian states mentioned.")

            # Crime pattern summary
            if crime_patterns.primary_crimes:
                top_crime = crime_patterns.primary_crimes[0][0].replace('_', ' ')
                summary_parts.append(f"Primary security concern relates to {top_crime} with "
                                     f"{len(crime_patterns.primary_crimes)} crime categories identified.")

            # Urgency assessment
            if sentiment_analysis.urgency_indicators:
                summary_parts.append(f"Document contains {len(sentiment_analysis.urgency_indicators)} "
                                     f"urgency indicators requiring attention.")

            # Text complexity
            if text_stats and text_stats.word_count > 1000:
                summary_parts.append(f"Comprehensive document with {text_stats.word_count:,} words "
                                     f"providing detailed intelligence data.")

            # Ensure minimum summary content
            if len(summary_parts) == 0:
                summary_parts.append("Intelligence analysis completed with standard processing.")

            final_summary = " ".join(summary_parts)

            # PROTECTION: Ensure summary is never empty
            if not final_summary or len(final_summary.strip()) < 10:
                final_summary = ("Intelligence document analysis completed successfully. "
                                 "Document processed for security intelligence extraction and threat assessment.")

            return final_summary

        except Exception as e:
            logger.error(f"Summary generation error: {str(e)}")
            # FALLBACK: Never return empty summary
            return ("Intelligence analysis completed with automated processing. "
                    "Document contains security-related information requiring analysis.")

    def _calculate_confidence_score(self, entities, sentiment_analysis,
                                    geographic_intel, temporal_intel) -> float:
        """Calculate overall analysis confidence score"""
        try:
            confidence_factors = []

            # Entity extraction confidence
            total_entities = sum(len(ent_list) for ent_list in entities.values())
            entity_confidence = min(total_entities / 10.0, 1.0)
            confidence_factors.append(entity_confidence)

            # Geographic confidence
            geo_confidence = min(geographic_intel.total_locations / 5.0, 1.0)
            confidence_factors.append(geo_confidence)

            # Temporal confidence
            temporal_elements = (len(temporal_intel.dates_mentioned) +
                                 len(temporal_intel.years_mentioned) +
                                 len(temporal_intel.months_mentioned))
            temporal_confidence = min(temporal_elements / 5.0, 1.0)
            confidence_factors.append(temporal_confidence)

            # Threat detection confidence
            threat_confidence = len(sentiment_analysis.urgency_indicators) / 10.0
            confidence_factors.append(min(threat_confidence, 1.0))

            # Calculate weighted average
            if confidence_factors:
                overall_confidence = sum(confidence_factors) / len(confidence_factors)
                return max(0.3, min(0.95, overall_confidence))  # Clamp between 30% and 95%
            else:
                return 0.5

        except Exception as e:
            logger.error(f"Confidence calculation error: {str(e)}")
            return 0.5


# Initialize analyzer
analyzer = IntelligenceAnalyzer()


# API Routes
@app.get("/")
async def root():
    """Health check endpoint"""
    return {
        "message": "Intelligence Document Analyzer API",
        "version": "3.0.0",
        "status": "operational",
        "timestamp": datetime.now().isoformat()
    }


@app.get("/health")
async def health_check():
    """Detailed health check"""
    return {
        "status": "healthy",
        "spacy_model": nlp is not None,
        "nltk_sentiment": sia is not None,
        "timestamp": datetime.now().isoformat()
    }


@app.post("/upload-document", response_model=AnalyzedDocument)
async def upload_document(file: UploadFile = File(...)):
    """Upload and analyze intelligence document"""
    try:
        # Validate file
        if not file.filename:
            raise HTTPException(status_code=400, detail="No file provided")

        # Read file content
        file_content = await file.read()
        file_size = len(file_content)

        if file_size == 0:
            raise HTTPException(status_code=400, detail="Empty file provided")

        if file_size > 50 * 1024 * 1024:  # 50MB limit
            raise HTTPException(status_code=400, detail="File too large (max 50MB)")

        # Extract text from file
        try:
            text_content = analyzer.extract_text_from_file(file_content, file.filename)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to extract text: {str(e)}")

        if not text_content.strip():
            raise HTTPException(status_code=400, detail="No readable text found in document")

        # Create metadata
        metadata = DocumentMetadata(
            filename=file.filename,
            file_type=file.filename.split('.')[-1].lower(),
            uploaded_at=datetime.now().isoformat(),
            file_size=file_size
        )

        # Perform analysis
        try:
            analysis = analyzer.analyze_document(text_content, metadata)
        except Exception as e:
            logger.error(f"Analysis error: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")

        # Create response
        document_id = str(uuid.uuid4())

        # PROTECTION: Ensure content is preserved (first 2000 chars to prevent memory issues)
        preserved_content = text_content[:2000] + "..." if len(text_content) > 2000 else text_content

        analyzed_document = AnalyzedDocument(
            id=document_id,
            content=preserved_content,
            metadata=metadata,
            analysis=analysis
        )

        logger.info(f"Successfully analyzed document: {file.filename}")
        return analyzed_document

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")


@app.get("/document-list")
async def get_document_list():
    """Get list of processed documents (mock data for demo)"""
    # In a real implementation, this would query a database
    mock_documents = [
        {
            "id": "doc_001",
            "filename": "intelligence_report_2025.pdf",
            "file_type": "pdf",
            "processed_at": "2025-06-21T10:30:00Z",
            "confidence_score": 0.92,
            "intelligence_summary": "High-priority intelligence report indicating increased security threats in Lagos region with multiple entities and locations identified."
        },
        {
            "id": "doc_002",
            "filename": "surveillance_data.docx",
            "file_type": "docx",
            "processed_at": "2025-06-21T09:15:00Z",
            "confidence_score": 0.78,
            "intelligence_summary": "Surveillance report documenting suspicious activities with moderate threat level and geographic intelligence data."
        },
        {
            "id": "doc_003",
            "filename": "incident_analysis.txt",
            "file_type": "txt",
            "processed_at": "2025-06-21T08:45:00Z",
            "confidence_score": 0.85,
            "intelligence_summary": "Incident analysis report with detailed entity extraction and crime pattern identification."
        }
    ]

    return {"documents": mock_documents}


@app.get("/document/{document_id}")
async def get_document_details(document_id: str):
    """Get detailed analysis for a specific document (mock for demo)"""
    # In a real implementation, this would query a database
    raise HTTPException(status_code=404, detail="Document not found")


@app.post("/query")
async def query_documents(query: dict):
    """Process AI query against document database"""
    try:
        user_query = query.get("query", "")

        if not user_query:
            raise HTTPException(status_code=400, detail="No query provided")

        # Enhanced AI response based on common intelligence queries
        query_lower = user_query.lower()

        if any(term in query_lower for term in ['threat', 'security', 'risk']):
            mock_response = f"""
**Security Threat Analysis: "{user_query}"**

Based on the analyzed intelligence documents, here are the key findings:

**Threat Assessment:**
• Current threat level: Medium to High
• Primary concerns: Armed banditry, cattle rustling, and organized crime
• Geographic hotspots: Zamfara, Katsina, Kaduna, and Plateau states

**Key Intelligence Findings:**
• 3 high-confidence intelligence reports analyzed
• Multiple security incidents documented across Northern Nigeria
• Coordinated criminal activities involving armed groups
• Increased incidents during specific time periods

**Geographic Analysis:**
• Activity concentration in North-West and North-Central zones
• Cross-border criminal movements detected
• Rural communities particularly vulnerable

**Recommendations:**
• Enhanced surveillance in identified geographic areas
• Increased security deployment in vulnerable zones
• Inter-agency coordination for comprehensive response
• Community engagement for improved intelligence gathering
"""
        elif any(term in query_lower for term in ['location', 'geographic', 'where']):
            mock_response = f"""
**Geographic Intelligence Analysis: "{user_query}"**

**Primary Areas of Concern:**
• Zamfara State: Highest incident concentration (15-20% of total cases)
• Katsina State: Secondary hotspot with cross-border activities
• Kaduna State: Strategic corridor with significant security challenges
• Plateau State: Emerging threat area requiring attention

**Geographic Patterns:**
• North-West zone: 40-50% of total security incidents
• Rural-urban crime corridors identified
• Seasonal migration patterns affecting security dynamics

**Infrastructure Impact:**
• Major highways: Abuja-Kaduna corridor particularly affected
• Border areas: Porous boundaries enabling criminal movement
• Remote communities: Limited security presence increases vulnerability
"""
        else:
            mock_response = f"""
**Intelligence Analysis Response: "{user_query}"**

Based on the available intelligence documents, here are the key findings:

**Document Analysis Results:**
• 3 documents contain relevant information for your query
• High-confidence matches found in security reports
• Cross-referenced intelligence from multiple sources

**Key Insights:**
• Criminal patterns indicate systematic organization
• Geographic concentration in specific Nigerian states
• Temporal analysis shows activity spikes during certain periods
• Entity analysis reveals connections between groups and locations

**Intelligence Summary:**
• Current threat assessment: Medium level
• Risk factors include resource availability and territorial expansion
• Recommended actions include enhanced monitoring and strategic deployment

**Data Sources:**
• Intelligence reports from Nigerian security agencies
• Geographic analysis covering 36 states plus FCT
• Temporal data spanning multiple months of 2020
"""

        return {"response": mock_response}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Query processing error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Query processing failed: {str(e)}")


if __name__ == "__main__":
    uvicorn.run(
        "main:app",
        host="0.0.0.0",
        port=8000,
        reload=True,
        log_level="info"
    )